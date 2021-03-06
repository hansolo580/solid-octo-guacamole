from docx import Document
from tkinter import Button, END, Entry, Label, OptionMenu, StringVar, Text, Tk
from tkinter import filedialog
from datetime import datetime
from datetime import date
from tkinter.messagebox import showerror
import os
from tkcalendar import Calendar, DateEntry

# Style
colorPrimary = '#AD2623'
colorSecondary = '#003352'
# End Style

# Window Rules
winMain = Tk()
winMain.title("Onesheet Generator")
winMain.geometry('1000x750')

winMain.grid_columnconfigure(0, weight=1)
winMain.grid_columnconfigure(1, weight=1)
winMain.grid_columnconfigure(2, weight=1)
winMain.grid_columnconfigure(3, weight=1)
winMain.grid_columnconfigure(4, weight=1)
winMain.grid_columnconfigure(5, weight=1)
winMain.grid_columnconfigure(6, weight=1)
winMain.grid_columnconfigure(7, weight=1)
winMain.grid_columnconfigure(8, weight=1)
winMain.grid_columnconfigure(9, weight=1)
winMain.grid_columnconfigure(10, weight=1)
winMain.grid_columnconfigure(11, weight=1)
winMain.grid_rowconfigure(0, weight=1)
winMain.grid_rowconfigure(1, weight=1)
winMain.grid_rowconfigure(2, weight=1)
winMain.grid_rowconfigure(3, weight=1)
winMain.grid_rowconfigure(4, weight=1)
winMain.grid_rowconfigure(5, weight=1)
winMain.grid_rowconfigure(6, weight=1)
winMain.grid_rowconfigure(7, weight=1)
winMain.grid_rowconfigure(8, weight=1)
winMain.grid_rowconfigure(9, weight=1)
winMain.grid_rowconfigure(10, weight=1)
winMain.grid_rowconfigure(11, weight=1)
winMain.grid_rowconfigure(12, weight=1)
winMain.grid_rowconfigure(13, weight=1)
winMain.grid_rowconfigure(14, weight=1)

Label(winMain, text="Job Title").grid(column=0, row=2)
jobTitle = Entry(winMain)
jobTitle.grid(column=1, row=2, columnspan=3)
jobTitle.insert(END, '')

Label(winMain, text="Client").grid(column=0, row=3)
client = Entry(winMain)
client.grid(column=1, row=3, columnspan=3)
client.insert(END, '')

Label(winMain, text="Order Date").grid(column=0, row=4)
orderDateSelector = DateEntry(winMain)
orderDateSelector.grid(column=1, row=4, columnspan=3)

Label(winMain, text="Start Date").grid(column=0, row=5)
startDateSelector = DateEntry(winMain)
startDateSelector.grid(column=1, row=5, columnspan=3)

Label(winMain, text="Pay Rate").grid(column=0, row=6)
payRate = Entry(winMain)
payRate.grid(column=1, row=6, columnspan=3)
payRate.insert(END, '')

Label(winMain, text="Heavy Lifter?").grid(column=0, row=7)
heavyLifter = StringVar(winMain)
heavyLifter.set(" ")
heavyLifterOptions = {"Yes": True, "No": False, " ": " "}
heavyLifterSelect = OptionMenu(winMain, heavyLifter, *heavyLifterOptions.keys())
heavyLifterSelect.grid(column=1, row=7, columnspan=3)

Label(winMain, text="Shift").grid(column=0, row=8)
shift = StringVar(winMain)
shift.set("TBD")
shiftOptions = {"First": 1, "Second": 2, "Third": 3, "TBD": "TBD"}
shiftSelect = OptionMenu(winMain, shift, *shiftOptions.keys())
shiftSelect.grid(column=1, row=8, columnspan=3)

Label(winMain, text="Hours").grid(column=0, row=9)
hours = Entry(winMain)
hours.grid(column=1, row=9, columnspan=3)
hours.insert(END, '')

Label(winMain, text="Location").grid(column=0, row=10)
location = Entry(winMain)
location.grid(column=1, row=10, columnspan=3)
location.insert(END, '')

Label(winMain, text="Supervisor").grid(column=0, row=11)
supervisor = Entry(winMain)
supervisor.grid(column=1, row=11, columnspan=3)
supervisor.insert(END, '')

Label(winMain, text="Number of Openings").grid(column=0, row=12)
openings = Entry(winMain)
openings.grid(column=1, row=12, columnspan=3)
openings.insert(END, '')

# SPACERS
Label(winMain, text="   ").grid(column=4, row=1)
Label(winMain, text="   ").grid(column=6, row=1)
Label(winMain, text="").grid(column=5, row=13)
Label(winMain, text="").grid(column=0, row=1)


# ENDSPACERS

def focus_next_widget(event):
    event.widget.tk_focusNext().focus()
    return "break"


Label(winMain, text="Job Description").grid(column=5, row=1)
jobDescription = Text(winMain, width=30, height=25)
jobDescription.grid(column=5, row=2, rowspan=6)
jobDescription.insert(END, '')
jobDescription.bind("<Tab>", focus_next_widget)

Label(winMain, text="Background Requirements").grid(column=5, row=8)
background = Text(winMain, width=30, height=5)
background.grid(column=5, row=9, rowspan=2)
background.insert(END, '')
background.bind("<Tab>", focus_next_widget)

Label(winMain, text="Education").grid(column=7, row=1)
education = Text(winMain, width=30, height=5)
education.grid(column=7, row=2, rowspan=2)
education.insert(END, '')
education.bind("<Tab>", focus_next_widget)

Label(winMain, text="Experience").grid(column=7, row=4)
experience = Text(winMain, width=30, height=5)
experience.grid(column=7, row=5, rowspan=2)
experience.insert(END, '')
experience.bind("<Tab>", focus_next_widget)

Label(winMain, text="Skills").grid(column=7, row=7)
skills = Text(winMain, width=30, height=5)
skills.grid(column=7, row=8, rowspan=2)
skills.insert(END, '')
skills.bind("<Tab>", focus_next_widget)

Label(winMain, text="Certifications").grid(column=7, row=10)
certifications = Text(winMain, width=30, height=5)
certifications.grid(column=7, row=11, rowspan=2)
certifications.insert(END, '')
certifications.bind("<Tab>", focus_next_widget)

saveLocation = os.path.join(str(os.getcwd()), 'Onesheets')
if not os.path.exists(saveLocation):
    os.makedirs(saveLocation)


def clk_generate_onesheet():
    job_title_value = jobTitle.get()
    client_value = client.get()
    order_date_value = str(orderDateSelector.get_date())
    start_date_value = str(startDateSelector.get_date())
    pay_rate_value = payRate.get()
    if heavyLifterOptions[heavyLifter.get()]:
        heavy_lifter_display = "Yes"
    elif not heavyLifterOptions[heavyLifter.get()]:
        heavy_lifter_display = "No"
    else:
        showerror(title="Error", message="You must select Yes or No for Heavy Lifter")
    shift_select_value = shiftOptions[shift.get()]
    hours_value = hours.get()
    location_value = location.get()
    supervisor_value = supervisor.get()
    openings_value = openings.get()
    background_value = background.get("1.0", 'end-1c')
    job_description_value = jobDescription.get('1.0', 'end-1c')
    education_value = education.get('1.0', 'end-1c')
    experience_value = experience.get('1.0', 'end-1c')
    skills_value = skills.get('1.0', 'end-1c')
    certifications_value = certifications.get('1.0', 'end-1c')

    onesheet = Document()

    onesheet.add_heading(job_title_value + " @ " + client_value, 0)

    # builds and display table
    display_table = onesheet.add_table(rows=15, cols=2)

    label_cells = display_table.columns[0].cells
    label_cells[0].text = 'Order Date'
    label_cells[1].text = 'Start Date'
    label_cells[2].text = 'Pay Rate'
    label_cells[3].text = 'Heavy Lifter'
    label_cells[4].text = 'Shift'
    label_cells[5].text = 'Hours'
    label_cells[6].text = 'Location'
    label_cells[7].text = 'Supervisor'
    label_cells[8].text = 'Openings'
    label_cells[9].text = 'Background Requirements'
    label_cells[10].text = 'Job Description'
    label_cells[11].text = 'Education Requirements'
    label_cells[12].text = 'Experience Requirements'
    label_cells[13].text = 'Skill Requirements'
    label_cells[14].text = 'Certification Requirements'

    data_cells = display_table.columns[1].cells
    data_cells[0].text = order_date_value
    data_cells[1].text = start_date_value
    data_cells[2].text = pay_rate_value
    data_cells[3].text = heavy_lifter_display
    data_cells[4].text = str(shift_select_value)
    data_cells[5].text = hours_value
    data_cells[6].text = location_value
    data_cells[7].text = supervisor_value
    data_cells[8].text = openings_value
    data_cells[9].text = background_value
    data_cells[10].text = job_description_value
    data_cells[11].text = education_value
    data_cells[12].text = experience_value
    data_cells[13].text = skills_value
    data_cells[14].text = certifications_value

    today = date.today()
    filename = str(client_value) + str(job_title_value) + str(today) + '.docx'
    onesheet.save(saveLocation + '\\' + filename)


btnCreateOnesheet = Button(winMain, text="Generate Onesheet", height=2, width=20, bg=colorPrimary, fg="white",
                           command=clk_generate_onesheet)
btnCreateOnesheet.grid(column=5, row=12)
btnCreateOnesheet.bind("<Return>", clk_generate_onesheet)

winMain.mainloop()
